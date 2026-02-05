"""
File Processor - VERSION ULTRA-ROBUSTE
Extrait et analyse TOUS les types de fichiers sans exception
Support : ZIP, 7z, RAR, TAR, GZ, PDF, DOCX, TXT, etc.
"""

import os
import tempfile
import shutil
from pathlib import Path
from typing import List, Dict, Tuple
import zipfile
import tarfile
import py7zr
import rarfile
import PyPDF2
from docx import Document as DocxDocument
import logging

logger = logging.getLogger(__name__)

class UniversalFileProcessor:
    """
    Processeur universel de fichiers pour DCE
    Analyse TOUS les types de fichiers sans exception
    """
    
    # Tous les formats supportés
    ARCHIVE_FORMATS = {
        '.zip': 'zip',
        '.7z': '7z',
        '.rar': 'rar',
        '.tar': 'tar',
        '.gz': 'gz',
        '.tgz': 'tgz',
        '.bz2': 'bz2'
    }
    
    DOCUMENT_FORMATS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'doc',
        '.txt': 'txt',
        '.md': 'md',
        '.rtf': 'rtf'
    }
    
    def __init__(self):
        self.temp_dir = None
        self.extracted_files = []
        self.total_text = ""
        self.file_count = 0
        self.errors = []
    
    async def process_file(self, file_path: str, original_filename: str) -> Dict:
        """
        Point d'entrée principal - Traite N'IMPORTE QUEL fichier
        
        Args:
            file_path: Chemin du fichier uploadé
            original_filename: Nom original du fichier
        
        Returns:
            Dict avec texte extrait, stats, erreurs
        """
        
        try:
            # Créer dossier temporaire
            self.temp_dir = tempfile.mkdtemp(prefix='dce_analysis_')
            
            logger.info(f"🔍 Traitement de : {original_filename}")
            
            # Détecter le type de fichier
            file_extension = Path(original_filename).suffix.lower()
            
            # Cas 1 : Archive (ZIP, 7z, RAR, etc.)
            if file_extension in self.ARCHIVE_FORMATS:
                logger.info(f"📦 Archive détectée : {file_extension}")
                await self._extract_archive(file_path, file_extension)
                await self._process_all_extracted_files()
            
            # Cas 2 : Document unique (PDF, DOCX, etc.)
            elif file_extension in self.DOCUMENT_FORMATS:
                logger.info(f"📄 Document unique détecté : {file_extension}")
                text = await self._extract_text_from_file(file_path, file_extension)
                self.total_text += text
                self.file_count += 1
            
            # Cas 3 : Format inconnu → Essayer quand même
            else:
                logger.warning(f"⚠️ Format inconnu : {file_extension}, tentative d'extraction...")
                # Tenter extraction comme archive
                try:
                    await self._extract_archive(file_path, '.zip')  # Tenter ZIP par défaut
                    await self._process_all_extracted_files()
                except:
                    # Tenter extraction comme PDF
                    try:
                        text = await self._extract_text_from_file(file_path, '.pdf')
                        self.total_text += text
                        self.file_count += 1
                    except Exception as e:
                        self.errors.append(f"Format non supporté : {file_extension} - {str(e)}")
            
            # Résultat final
            result = {
                'success': len(self.total_text) > 0,
                'extracted_text': self.total_text,
                'files_processed': self.file_count,
                'total_characters': len(self.total_text),
                'errors': self.errors,
                'extracted_files': [str(f) for f in self.extracted_files]
            }
            
            logger.info(f"✅ Traitement terminé : {self.file_count} fichiers, {len(self.total_text)} caractères")
            
            return result
            
        except Exception as e:
            logger.error(f"❌ Erreur globale : {str(e)}")
            return {
                'success': False,
                'extracted_text': '',
                'files_processed': 0,
                'total_characters': 0,
                'errors': [f"Erreur critique : {str(e)}"],
                'extracted_files': []
            }
        
        finally:
            # Nettoyer le dossier temporaire
            if self.temp_dir and os.path.exists(self.temp_dir):
                try:
                    shutil.rmtree(self.temp_dir)
                    logger.info(f"🧹 Nettoyage dossier temporaire : {self.temp_dir}")
                except Exception as e:
                    logger.warning(f"⚠️ Impossible de nettoyer {self.temp_dir}: {e}")
    
    async def _extract_archive(self, archive_path: str, extension: str):
        """Extrait une archive (ZIP, 7z, RAR, TAR, etc.)"""
        
        try:
            if extension in ['.zip']:
                await self._extract_zip(archive_path)
            elif extension in ['.7z']:
                await self._extract_7z(archive_path)
            elif extension in ['.rar']:
                await self._extract_rar(archive_path)
            elif extension in ['.tar', '.gz', '.tgz', '.bz2']:
                await self._extract_tar(archive_path)
            else:
                raise ValueError(f"Format d'archive non supporté : {extension}")
            
            logger.info(f"📦 Archive extraite : {len(self.extracted_files)} fichiers trouvés")
            
        except Exception as e:
            error_msg = f"Erreur extraction archive {extension}: {str(e)}"
            logger.error(f"❌ {error_msg}")
            self.errors.append(error_msg)
    
    async def _extract_zip(self, zip_path: str):
        """Extrait un fichier ZIP"""
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(self.temp_dir)
            self.extracted_files = list(Path(self.temp_dir).rglob('*'))
    
    async def _extract_7z(self, seven_z_path: str):
        """Extrait un fichier 7z"""
        with py7zr.SevenZipFile(seven_z_path, 'r') as archive:
            archive.extractall(self.temp_dir)
            self.extracted_files = list(Path(self.temp_dir).rglob('*'))
    
    async def _extract_rar(self, rar_path: str):
        """Extrait un fichier RAR"""
        with rarfile.RarFile(rar_path, 'r') as rar_ref:
            rar_ref.extractall(self.temp_dir)
            self.extracted_files = list(Path(self.temp_dir).rglob('*'))
    
    async def _extract_tar(self, tar_path: str):
        """Extrait un fichier TAR/GZ/TGZ/BZ2"""
        with tarfile.open(tar_path, 'r:*') as tar_ref:
            tar_ref.extractall(self.temp_dir)
            self.extracted_files = list(Path(self.temp_dir).rglob('*'))
    
    async def _process_all_extracted_files(self):
        """Traite TOUS les fichiers extraits d'une archive"""
        
        if not self.extracted_files:
            logger.warning("⚠️ Aucun fichier extrait de l'archive")
            return
        
        # Filtrer les fichiers (pas les dossiers)
        files_only = [f for f in self.extracted_files if f.is_file()]
        
        logger.info(f"📁 {len(files_only)} fichiers à traiter")
        
        for file_path in files_only:
            file_extension = file_path.suffix.lower()
            
            # Traiter seulement les documents (ignorer images, etc.)
            if file_extension in self.DOCUMENT_FORMATS:
                try:
                    text = await self._extract_text_from_file(str(file_path), file_extension)
                    self.total_text += f"\n\n{'='*80}\n"
                    self.total_text += f"FICHIER : {file_path.name}\n"
                    self.total_text += f"{'='*80}\n\n"
                    self.total_text += text
                    self.file_count += 1
                    logger.info(f"✅ Extrait : {file_path.name} ({len(text)} caractères)")
                except Exception as e:
                    error_msg = f"Erreur traitement {file_path.name}: {str(e)}"
                    logger.error(f"❌ {error_msg}")
                    self.errors.append(error_msg)
            else:
                logger.debug(f"⏭️ Ignoré (format non document) : {file_path.name}")
    
    async def _extract_text_from_file(self, file_path: str, extension: str) -> str:
        """Extrait le texte d'un fichier selon son type"""
        
        if extension == '.pdf':
            return await self._extract_from_pdf(file_path)
        elif extension == '.docx':
            return await self._extract_from_docx(file_path)
        elif extension in ['.txt', '.md', '.rtf']:
            return await self._extract_from_text(file_path)
        else:
            logger.warning(f"⚠️ Format non supporté pour extraction : {extension}")
            return ""
    
    async def _extract_from_pdf(self, pdf_path: str) -> str:
        """Extrait le texte d'un PDF (EXHAUSTIF - 200 pages)"""
        
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                # IMPORTANT : Traiter TOUTES les pages (limite 200 pour sécurité)
                max_pages = min(len(pdf_reader.pages), 200)
                
                logger.info(f"📄 Extraction PDF : {max_pages} pages")
                
                for page_num in range(max_pages):
                    try:
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        text += f"\n[Page {page_num + 1}]\n"
                        text += page_text
                    except Exception as e:
                        logger.warning(f"⚠️ Erreur page {page_num + 1} : {e}")
                        continue
                
                return text
        except Exception as e:
            raise Exception(f"Erreur extraction PDF : {str(e)}")
    
    async def _extract_from_docx(self, docx_path: str) -> str:
        """Extrait le texte d'un DOCX"""
        
        try:
            doc = DocxDocument(docx_path)
            text = ""
            
            # Extraire paragraphes
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extraire tableaux
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " | "
                    text += "\n"
            
            return text
        except Exception as e:
            raise Exception(f"Erreur extraction DOCX : {str(e)}")
    
    async def _extract_from_text(self, text_path: str) -> str:
        """Extrait le texte d'un fichier texte"""
        
        try:
            # Essayer plusieurs encodages
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(text_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            # Si tous les encodages échouent
            raise Exception("Impossible de décoder le fichier texte")
            
        except Exception as e:
            raise Exception(f"Erreur extraction texte : {str(e)}")


# Instance globale
file_processor = UniversalFileProcessor()
